#extracting image names
from docx import *
import re
from pathlib import Path
import os

CUR_BASE_DIR = os.path.join("/home/kushal/Desktop/MTP_project/Django_MTP" ,"mysite", "media2")

#data_folder = Path(r"")
#file_to_open = data_folder / "newfile2.txt"

#f = open(file_to_open)
#p = f.read()
#print(p)
document = Document(os.path.join(CUR_BASE_DIR, 'Sample_1.docx'))
#document = Document('Sample_1.docx')
lists = []
guestFileName = os.path.join(CUR_BASE_DIR, 'Image_extraction.csv')
guestFile = open(guestFileName,"wb")
guestFile.write("Image Id".encode("utf-8")+",".encode("utf-8"))
guestFile.write("Image Names".encode("utf-8")+"\n".encode("utf-8"))
# guestFile.write("s1, s2 \n".encode("utf-8"))

for para in document.paragraphs:
    f2=0
    f1=0
    if para.text=="":
        continue
    for run in para.runs:
        if run.bold :
            f1=1
            break
        else:
            f2=1
    #print(para.text)
    if len(para.text)>=1:
       # print(para.text)
        str = para.text.split(' ')[0]
        x = re.findall("Figure", str)
       # if f2!=1:
          #  print(str,x)
    if f2!=1 and x:
        lists.append(para.text.split('. ')[0])
        #print(para.text.split('. ')[0])
        #print(para.text.split('. ')[1])
        guestFile = open(guestFileName,"ab")
        s1 = para.text.split('. ')[0]
        s2 = para.text.split('. ')[1]
        s = (s1.encode("utf-8")+",".encode("utf-8")+s2.encode("utf-8"))
        guestFile.write(s)
       # print(s1,s2)
        guestFile.write("\n".encode("utf-8"))
        guestFile.close()
        #print(x,y)
        f2=0


# Added encode

import  pandas as pd
csv_file_name = os.path.join(CUR_BASE_DIR, 'Image_extraction.csv')
df = pd.read_csv(csv_file_name)
list_of_dict= df.to_dict('records')
print(list_of_dict)
print("Printing individual dict :: \n")

#l = [ 'ram', 'shyam ','mohan']
#for index, item in enumerate(l):
#    print(index, item)
for index, temp_list in enumerate(list_of_dict):
    for key in temp_list:
        print(temp_list[key], end= '    ')
    print()

