from docx.api import Document
from pathlib import Path

data_folder = Path(r"")
file_to_open = data_folder / "newfile2.txt"

f = open(file_to_open)
p = f.read()
print(p)
document = Document(p)
f = 0
for t in document.tables:
    table = t
    #print(len(document.tables))
    data = []
    guestFile = open("Table_Extraction.csv","w")
    guestFile.write("Table Id"+",")
    guestFile.write("Table Name"+"\n")
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        keys = tuple(text)
        print(keys[0])
        for s in keys[0].split(' '):
            if f==2:
                guestFile.write(",")
            guestFile.write(s.replace('.',' '))
            f = f+1
        break;
guestFile.write("\n")
guestFile.close()

#extracting image names
from docx import *
import re
from pathlib import Path

document = Document(p)
lists = []
f = 0
guestFile = open("Table_Extraction.csv","a")
for para in document.paragraphs:
    if f==1:
        break
    s = para.text.split(' ')
    #print(para.text)
    for run in para.runs:
        if run.bold:
            if f==1:
                break
          #  print(para.text)
            for s1 in s:
                if s1=="Table.":
                    for s in para.text.split(' '):
                        if f==2:
                            guestFile.write(",")
                        guestFile.write(s.replace('.',' '))
                       # print(s)
                        f = f+1
        break;
 #   if para.text=="Table":
  #      print(para.text)
guestFile.close()
