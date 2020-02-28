from docx import Document

document = Document('/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/latex.docx')
document.save('/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/latex.docx')

import re
dictlabel={}
c=0
for paragraph in document.paragraphs:
    if '\label{' in paragraph.text:
        c = c+1
     #   print(paragraph.text)
        text = paragraph.text
        m = re.findall('{(.+?)}', text)
        for k in m:
       # print(m.group(1))
            dictlabel[k.lower().replace(" ","")] = c      
print(dictlabel)     

import re
filepath = '/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/latex.txt'
dictlabel={}
c=0
with open(filepath,errors='ignore') as fp:
       line = fp.readline()
       while line:
          #  print(line)
            if 'section{' in line and '\label{' in line: 
                    m = re.findall('{(.+?)}', line)
                    dictlabel[m[1].lower().replace(" ","")] = c    
                    c = c+1
            line = fp.readline()
print(dictlabel)                

import itertools 

# List initialization 
list1 =dictlabel.keys() 
list2 =dictlabel.keys()  
  
# using itertools 
temp = list(itertools.product(list1, list1)) 
  
# output list initialization 
out = {} 
  
# iteration 
for elem in temp: 
    #if elem[0]!= elem[1]: 
        out[elem]=0 
#print(out['relwork','proposedalgorithm']) 

import re
section_name=""
#print(out['relwork','proposedalgorithm']) 
for paragraph in document.paragraphs :
    if '\label{' in paragraph.text:
        text = paragraph.text
        m = re.search('{(.+?)}', text)
        section_name =  m.group(1).lower().replace(" ","")
    if '\\ref{' in paragraph.text:
       # print(paragraph.text)
        text = paragraph.text
        m = re.findall('ref{(.+?)}', text)
        for k in m:
          #  print(k)
            section1 = k.lower().replace(" ","")
            tuples = k.lower().replace(" ",""),section_name
            if tuples in out.keys():
                out[section1,section_name] = out[section1,section_name]+1
           # line = line+1
          #  streetno[line] = k
                print(section1+" "+section_name)
                print(out[section1,section_name])
           # print(section_name)
        
import operator
from operator import itemgetter
guestFile = open("/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/Section_Section_Correlation_sorted.csv","w")
guestFile.close()
out1 = dict(sorted(out.items(),key=operator.itemgetter(1),reverse=True))
sorted(out.items(),key=operator.itemgetter(1),reverse=True)
#print(out1)
for entries in out1 :
    guestFile = open("/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/Section_Section_Correlation_sorted.csv","a")
    s = entries[0]+","+entries[1]+","+str(out.get(entries))
    guestFile.write(s)
    print(s)
    guestFile.write("\n")
    guestFile.close()

	
