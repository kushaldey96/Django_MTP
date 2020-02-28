#name of the sections
from docx import *
import re
document = Document('/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/Sample_2.docx')
guestFile = open("/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/References.csv","wb")
f=0
f2=0
f3 =0
f4 =0
flag = 0
guestFile.write("Authors".encode("utf-8"))
guestFile.write(",".encode("utf-8"))
guestFile.write("Year".encode("utf-8"))
guestFile.write(",".encode("utf-8"))
guestFile.write("paper name".encode("utf-8"))
for para in document.paragraphs:
    if para.text == "References":
        f=1
        continue;
    if para.text=="Figure captions":
        break
    if len(para.text)>=1 and f==1: 
       # print(para.text)
        s = para.text.split(',')
        guestFile.write("\n".encode("utf-8"))
        flag = 0
        for i in s:
            if len(i)>=20:
                flag  =  1
                i1 = i.split('.')
                for i2 in i1:
                    if len(i2)>=1:
                        for i3 in i2:
                            if i3>='0'and i3<='9':
                                f3 = 1
                                f4 = 1
                                break    
                    if f3==1:
                        guestFile.write(",".encode("utf-8")+i2.replace("(","").replace(")","").encode("utf-8")+" ".encode("utf-8"))
                        f3=0
                        guestFile.write(",".encode("utf-8")+" ".encode("utf-8"))
                    else:   
                        guestFile.write(i2.replace("(","").replace(")","").encode("utf-8")+" ".encode("utf-8"))
           # else:
            #    if f4==1:
             #       guestFile.write(",".encode("utf-8")+i.encode("utf-8"))
            else:    
                if flag == 1:
                    guestFile.write(",".encode("utf-8"))
                guestFile.write(i.encode("utf-8")+"  ".encode("utf-8")+" ".encode("utf-8"))
guestFile.close()                

from docx import *
import re
document = Document('/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/Sample_2.docx')
guestFile = open("/home/kushal/Desktop/MTP_project/Django_code/mysite/media2/References.csv","wb")
guestFile.write("Authors".encode("utf-8"))
guestFile.write(",".encode("utf-8"))
guestFile.write("Year".encode("utf-8"))
guestFile.write(",".encode("utf-8"))
guestFile.write("paper name".encode("utf-8"))
guestFile.write(",".encode("utf-8"))
guestFile.write("journal name".encode("utf-8"))
f = 0
flag=0
flag2=0
brk = 0
for p in document.paragraphs:
    if p.text=="References":
        f = 1
        continue
    if p.text=="Figure captions":
        break    
    if f==1: 
        guestFile.write("\n".encode("utf-8"))
        flag2 =0
        f2=0
        for run in p.runs:
            if run.italic:
               # guestFile.write(",".encode("utf-8")+run.text.replace(",","").encode("utf-8"))
               # f2 = f2+1
               # if f2==2:
               #     guestFile.write(",".encode("utf-8"))
                guestFile.write(",".encode("utf-8"))  
                for s6 in run.text.split(","):    
                  #  print(s6)
                    guestFile.write(" ".encode("utf-8")+s6.encode("utf-8"))
                guestFile.write(",".encode("utf-8"))
                break;
            else:
                s = run.text.split(',')
                flag=0
                for s1 in s:
                    if len(s1)>=30:
                        guestFile.write(",".encode("utf-8"))
                    #    print(s1)
                        s2 = s1.split(" ")
                        for s3 in s2:
                            for c in s3:
                                if c>='0' and c<='9':
                                    flag = 1;
                            if flag2==1:        
                                guestFile.write(" ".encode("utf-8")+s3.encode("utf-8"))
                            
                            if flag==1:
                                s4 = ""
                                for c in s3:
                                    if c>='0' and c<='9':
                                        s4 = s4+c
                             #   print(s4,s3)        
                                guestFile.write(s4.encode("utf-8"))
                                flag2 = 1
                        #    else:    
                        #        print("hh",s3)
                            if flag == 1:
                                guestFile.write(",".encode("utf-8"))
                                flag =  0
                    else:
                        guestFile.write(" ".encode("utf-8")+s1.encode("utf-8"))
                       # print(s1)
guestFile.close()                        
